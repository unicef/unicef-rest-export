import os
from io import BytesIO, StringIO
from tempfile import mkstemp

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Table, TableStyle
from reportlab.platypus.doctemplate import LayoutError
from rest_framework import status
from rest_framework.renderers import BaseRenderer, TemplateHTMLRenderer
from rest_framework_csv.renderers import CSVRenderer
from tablib import Dataset

RESPONSE_ERROR = (
    "Response data is a %s, not a Dataset! "
    "Did you extend ExportMixin?"
)
PDF_COLUMNS_PER_PAGE = 9


class ExportBaseRenderer(BaseRenderer):
    """Renders Datasets using their built in export implementation.
    Only works with serializers that return Datasets as their data object.
    Uses a StringIO to capture the output of dataset.export('[format]')
    """
    def render(self, data, accepted_media_type=None, renderer_context=None):
        if 'response' in renderer_context:
            status_code = renderer_context['response'].status_code
            if not status.is_success(status_code):
                return "Error: %s" % data.get('detail', status_code)

        if not isinstance(data, Dataset):
            raise Exception(
                RESPONSE_ERROR % type(data).__name__
            )

        self.init_output()
        args = self.get_export_args(data)
        kwargs = self.get_export_kwargs(data, renderer_context)
        self.render_dataset(data, *args, **kwargs)
        return self.get_output()

    def render_dataset(self, data, *args, **kwargs):
        self.output.write(data.export(self.format))

    def init_output(self):
        self.output = StringIO()

    def get_output(self):
        return self.output.getvalue()

    def get_export_args(self, data):
        return [self.output]

    def get_export_kwargs(self, data, renderer_context):
        return {}


class ExportFileRenderer(ExportBaseRenderer):
    """Renderer for output formats that absolutely must use a file
    (i.e. Excel)
    """
    def init_output(self):
        file, filename = mkstemp(suffix='.' + self.format)
        self.filename = filename
        os.close(file)

    def render_dataset(self, data, *args, **kwargs):
        with open(self.filename, "wb") as fp:
            fp.write(data.export(self.format))

    def get_export_args(self, data):
        return [self.filename]

    def get_output(self):
        file = open(self.filename, 'rb')
        result = file.read()
        file.close()
        os.unlink(self.filename)
        return result


class ExportHTMLRenderer(TemplateHTMLRenderer, ExportBaseRenderer):
    media_type = "text/html"
    format = "html"

    def render(self, data, accepted_media_type=None, renderer_context=None):
        table = ExportBaseRenderer.render(
            self, data, accepted_media_type, renderer_context,
        )

        return TemplateHTMLRenderer.render(
            self, {'table': table}, accepted_media_type, renderer_context,
        )

    def get_template_context(self, data, renderer_context):
        view = renderer_context['view']
        request = renderer_context['request']

        data['name'] = view.get_view_name()
        data['description'] = view.get_view_description(html=True)
        data['url'] = request.path.replace('.html', '')
        full_path = request.get_full_path()
        if '?' in full_path:
            data['url_params'] = full_path[full_path.index('?'):]
        data['available_formats'] = [
            cls.format for cls in view.renderer_classes
            if cls.format != 'html'
        ]

        if hasattr(view, 'get_template_context'):
            data.update(view.get_template_context(data))

        return data

    def get_export_kwargs(self, data, renderer_context):
        return {
            'classes': 'ui-table table-stripe',
            'na_rep': '',
        }


class ExportCSVRenderer(ExportBaseRenderer):
    """
    Renders data frame as CSV
    """
    media_type = "text/csv"
    format = "csv"

    def get_export_kwargs(self, data, renderer_context):
        return {'encoding': self.charset}


class ExportJSONRenderer(ExportBaseRenderer):
    """Renders dataset as JSON"""
    media_type = "application/json"
    format = "json"

    date_format_choices = {'epoch', 'iso'}
    default_date_format = 'iso'

    def get_export_kwargs(self, data, renderer_context):
        request = renderer_context['request']

        date_format = request.GET.get('date_format', '')
        if date_format not in self.date_format_choices:
            date_format = self.default_date_format

        return {
            'date_format': date_format,
        }


class ExportOpenXMLRenderer(ExportFileRenderer):
    """Renders dataset as Excel (.xlsx)"""
    media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"  # noqa
    format = "xlsx"


class ExportExcelRenderer(ExportFileRenderer):
    """Renders dataset as Excel (.xls)"""
    media_type = "application/vnd.ms-excel"
    format = "xls"


class ExportPDFTableRenderer(ExportFileRenderer):
    """Renders dataset as PDF (.pdf) in table format"""
    media_type = "application/pdf"
    format = "pdf_table"

    def export_set(self, formatted, headers, columns_per_page):
        stream = BytesIO()
        doc = SimpleDocTemplate(stream, pagesize=landscape(letter))
        styles = getSampleStyleSheet()
        styleCell = styles["Normal"]
        styleCell.fontSize = 7
        elements = []

        if headers:
            # slice the data into a set number of columns
            # in order to fit on the page
            for start in range(0, len(headers), columns_per_page):
                end = start + columns_per_page
                data = [['Row'] + [
                    item if item is not None else ''
                    for item in headers[start:end]
                ]]

                row_num = 1
                for row in formatted:
                    d = [
                        Paragraph(str(v), styleCell)
                        for _, v in row.items()
                    ]
                    data.append([row_num] + d[start:end])
                    row_num += 1

                t = Table(data)
                t.setStyle(TableStyle([
                    # action/format, from-cell, to-cell, format
                    ('VALIGN', (0, 0), (-1, -1), "TOP"),
                    ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
                    ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
                ]))
                elements.append(t)
                elements.append(PageBreak())

        doc.build(elements)
        return stream.getvalue()

    def export_failure(self):
        stream = BytesIO()
        doc = SimpleDocTemplate(stream, pagesize=landscape(letter))
        styles = getSampleStyleSheet()
        styleCell = styles["Normal"]
        styleCell.fontSize = 7
        elements = [
            Paragraph(
                (
                    "Data not able to be formatted for PDF, "
                    "please try another format."
                ),
                styleCell
            )
        ]
        doc.build(elements)
        return stream.getvalue()

    def render_dataset(self, data, *args, **kwargs):
        formatted = data._package()
        headers = data.headers
        with open(self.filename, "wb") as fp:
            columns_per_page = PDF_COLUMNS_PER_PAGE
            success = False
            while columns_per_page > 1:
                try:
                    fp.write(
                        self.export_set(formatted, headers, columns_per_page)
                    )
                except LayoutError:
                    columns_per_page -= 1
                else:
                    success = True
                    break
            if not success:
                fp.write(self.export_failure())


class ExportPDFRenderer(ExportFileRenderer):
    """Renders dataset as PDF (.pdf)"""
    media_type = "application/pdf"
    format = "pdf"

    def export_set(self, formatted):
        stream = BytesIO()
        doc = SimpleDocTemplate(stream)
        styles = getSampleStyleSheet()
        styleParagraph = styles["Normal"]
        styleParagraph.fontSize = 7
        elements = []

        for row in formatted:
            if row:
                for k, val in row.items():
                    p = Paragraph(f"<b>{k}:</b> {val}", styleParagraph)
                    elements.append(p)
                elements.append(PageBreak())

        doc.build(elements)
        return stream.getvalue()

    def render_dataset(self, data, *args, **kwargs):
        formatted = data._package()
        with open(self.filename, "wb") as fp:
            fp.write(self.export_set(formatted))


class ExportDocxBaseRenderer(ExportFileRenderer):
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ExportDocxRenderer(ExportDocxBaseRenderer):
    """Renders dataset as Doc (.docx)"""
    format = "docx"

    def export_set(self, formatted):
        stream = BytesIO()
        doc = Document()

        for row in formatted:
            if row:
                for k, val in row.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{k}:").bold = True
                    p.add_run(f"{val}")
                doc.add_page_break()

        doc.save(stream)
        return stream.getvalue()

    def render_dataset(self, data, *args, **kwargs):
        formatted = data._package()
        with open(self.filename, "wb") as fp:
            fp.write(self.export_set(formatted))


class ExportDocxTableRenderer(ExportDocxBaseRenderer):
    """Renders dataset as Doc (.docx) in table format"""
    format = "docx_table"

    def export_set(self, formatted, headers):
        stream = BytesIO()
        doc = Document()

        if not headers:
            doc.add_paragraph("No data provided.")
        else:
            table = doc.add_table(rows=1, cols=len(headers))

            # set heading text
            header_cells = table.rows[0].cells
            for i, heading in enumerate(headers):
                header_cells[i].text = heading

            # set data
            for record in formatted:
                row = table.add_row().cells
                for i, data in enumerate(record):
                    row[i].text = str(data)

        doc.save(stream)
        return stream.getvalue()

    def render_dataset(self, data, *args, **kwargs):
        formatted = data._package()
        headers = data.headers
        with open(self.filename, "wb") as fp:
            fp.write(self.export_set(formatted, headers))


class FriendlyCSVRenderer(CSVRenderer):
    def flatten_item(self, item):
        if isinstance(item, bool):
            return {'': {True: 'Yes', False: ''}[item]}
        return super().flatten_item(item)
